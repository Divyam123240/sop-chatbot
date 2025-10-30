import os
import io
import logging
import traceback
from typing import List
from datetime import datetime

import requests
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from sqlalchemy import create_engine, Column, Integer, Text, DateTime, text
from sqlalchemy.orm import sessionmaker, declarative_base
from sqlalchemy.dialects.postgresql import JSONB
from pgvector.sqlalchemy import Vector

from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

from PyPDF2 import PdfReader
import docx
from PIL import Image

# ----------------------------
# Load environment variables
# ----------------------------
load_dotenv()

DB_USER = os.getenv("DB_USER", "dev_admin")
DB_PASSWORD = os.getenv("DB_PASSWORD", "password")
DB_HOST = os.getenv("DB_HOST", "localhost")
DB_PORT = os.getenv("DB_PORT", "5432")
DB_NAME = os.getenv("DB_NAME", "sop_dev_db")

DATABASE_URL = (
    f"postgresql+psycopg2://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}?sslmode=require"
)

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "llama3:latest")

# ----------------------------
# Database setup
# ----------------------------
engine = create_engine(DATABASE_URL, echo=True, pool_pre_ping=True)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    content = Column(Text, nullable=False)
    doc_metadata = Column(JSONB, nullable=True)
    embedding = Column(Vector(768))


class ChatHistory(Base):
    __tablename__ = "chat_history"
    id = Column(Integer, primary_key=True, index=True)
    user_query = Column(Text, nullable=False)
    answer = Column(Text, nullable=False)
    timestamp = Column(DateTime, default=datetime.utcnow)


Base.metadata.create_all(bind=engine)

# ----------------------------
# Embedding model
# ----------------------------
EMBED_MODEL = SentenceTransformer("all-mpnet-base-v2")


def embed_texts(texts: List[str]) -> List[List[float]]:
    embeddings = EMBED_MODEL.encode(texts, convert_to_numpy=True, show_progress_bar=False)
    return [emb.tolist() for emb in embeddings]

# ----------------------------
# FastAPI setup
# ----------------------------
app = FastAPI(title="SOP Chatbot (Ollama + PostgreSQL + pgvector)")

# logger for diagnostics
logger = logging.getLogger("sop_upload")
logger.setLevel(logging.INFO)

# Enable CORS for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For dev; restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ----------------------------
# Request/Response models
# ----------------------------
class ChatRequest(BaseModel):
    query: str

# ----------------------------
# Root route & Health check
# ----------------------------
@app.get("/")
def root():
    return {"message": "SOP Chatbot backend is running. Use /api/chat or /api/upload."}

@app.get("/health")
def health_check():
    db = SessionLocal()
    try:
        db.execute(text("SELECT 1"))
        return {"status": "ok", "db": "connected"}
    except Exception as e:
        return {"status": "error", "db": "disconnected", "error": str(e)}
    finally:
        db.close()

# ----------------------------
# File Upload Endpoint
# ----------------------------
@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    # If filename has no extension, try to infer from UploadFile.content_type
    if not ext:
        ct = getattr(file, "content_type", "") or ""
        ct = ct.lower()
        if ct.startswith("image/"):
            ext = "." + ct.split("/")[-1]
        elif "pdf" in ct:
            ext = ".pdf"
        elif "wordprocessingml" in ct or "officedocument.wordprocessingml" in ct:
            ext = ".docx"
        elif "msword" in ct:
            ext = ".doc"
    extracted_text = ""

    file_bytes = await file.read()
    logger.info(f"Upload received: filename={filename!r} ext={ext!r} content_type={getattr(file, 'content_type', None)!r} size={len(file_bytes)} bytes")
    try:
        if ext in [".txt", ".md"]:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        elif ext == ".pdf":
            pdf = PdfReader(io.BytesIO(file_bytes))
            for page in pdf.pages:
                extracted_text += page.extract_text() or ""
        elif ext == ".docx":
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
        elif ext == ".doc":
            # Try to extract text from legacy .doc files using textract if available.
            # textract is optional because it requires native dependencies on some systems.
            try:
                import textract
                import tempfile

                # textract expects a filename on disk, so write to a temp file first
                with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
                    tf.write(file_bytes)
                    temp_path = tf.name

                try:
                    extracted = textract.process(temp_path, extension="doc")
                    extracted_text = extracted.decode("utf-8", errors="ignore")
                finally:
                    # remove temp file if it exists
                    try:
                        os.unlink(temp_path)
                    except Exception:
                        pass
            except Exception as tex_err:
                # If textract isn't available or fails, provide a helpful error message.
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Unsupported .doc processing: textract not available or failed. "
                        "Please convert the .doc file to .docx or PDF and try again, or install textract with its dependencies. "
                        f"(textract error: {str(tex_err)})"
                    ),
                )
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            image = Image.open(io.BytesIO(file_bytes))
            extracted_text = f"Image file: {filename}, format: {image.format}, size: {image.size}"
        else:
            # Provide a clearer error message listing supported types
            supported = ".txt, .md, .pdf, .docx, .png, .jpg, .jpeg, .bmp, .gif, .webp"
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type '{ext}' for file '{filename}'. Supported types: {supported}",
            )
    except Exception as e:
        # log full traceback to server logs for debugging
        tb = traceback.format_exc()
        logger.error(f"Error processing file {filename}: {str(e)}\n{tb}")
        raise HTTPException(status_code=400, detail=f"File processing error: {str(e)}")

    # If extraction produced no text, try some optional fallbacks (pdfplumber, pytesseract, soffice)
    if not extracted_text or not extracted_text.strip():
        # PDF fallback: try pdfplumber if available
        if ext == ".pdf":
            try:
                import pdfplumber

                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        extracted_text += (page.extract_text() or "")
            except Exception:
                pass

        # Image OCR fallback: try pytesseract if available
        if (not extracted_text or not extracted_text.strip()) and ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            try:
                import pytesseract

                image = Image.open(io.BytesIO(file_bytes))
                extracted_text = pytesseract.image_to_string(image)
            except Exception:
                pass

        # Legacy .doc fallback: try LibreOffice (soffice) conversion to docx if available
        if (not extracted_text or not extracted_text.strip()) and ext == ".doc":
            try:
                import shutil
                import subprocess
                import tempfile

                soffice = shutil.which("soffice") or shutil.which("libreoffice")
                if soffice:
                    with tempfile.TemporaryDirectory() as td:
                        in_path = os.path.join(td, filename)
                        with open(in_path, "wb") as f:
                            f.write(file_bytes)
                        # Convert to docx
                        subprocess.run([
                            soffice,
                            "--headless",
                            "--convert-to",
                            "docx",
                            "--outdir",
                            td,
                            in_path,
                        ], check=True)
                        base = os.path.splitext(filename)[0]
                        converted = os.path.join(td, base + ".docx")
                        if os.path.exists(converted):
                            doc = docx.Document(converted)
                            for para in doc.paragraphs:
                                extracted_text += para.text + "\n"
            except Exception:
                pass

    # Split text into chunks for embedding
    if ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
        chunks = [extracted_text]
    else:
        chunks = [p.strip() for p in extracted_text.split("\n\n") if p.strip()]

    if not chunks:
        logger.info(f"No text content extracted for {filename}. ext={ext} extracted_text_len={len(extracted_text)}")
        # Provide more actionable guidance depending on file type
        if ext == ".pdf":
            detail = (
                "No text content found in PDF. This may be a scanned/image PDF or a PDF with non-extractable text. "
                "Install optional tools to improve extraction: `pip install pdfplumber` for better PDF parsing, or "
                "install Tesseract OCR and `pip install pytesseract` to OCR image PDFs. Alternatively convert the PDF to a text-based PDF or to DOCX and re-upload."
            )
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"]:
            detail = (
                "No text found in image. If this is a scanned image containing text, install Tesseract OCR (system package) and `pip install pytesseract` to enable OCR."
            )
        else:
            detail = "No text content found in file."
        raise HTTPException(status_code=400, detail=detail)

    embeds = embed_texts(chunks)

    db = SessionLocal()
    try:
        for c, e in zip(chunks, embeds):
            doc = Document(content=c, doc_metadata={"source": filename}, embedding=e)
            db.add(doc)
        db.commit()
    finally:
        db.close()

    return {"status": "ok", "filetype": ext, "chunks": len(chunks)}

# ----------------------------
# Chat Endpoint
# ----------------------------
@app.post("/api/chat")
async def chat(request: ChatRequest):
    try:
        query = request.query
        db = SessionLocal()
        
        # Get embeddings for the query
        query_embedding = embed_texts([query])[0]
        
        # Search similar documents
        emb_str = ",".join(str(x) for x in query_embedding)
        sql = text(
            f"SELECT content FROM documents "
            f"ORDER BY embedding <-> ARRAY[{emb_str}]::vector LIMIT 3"
        )
        results = db.execute(sql).fetchall()
        
        # Prepare context from similar documents
        context = "\n".join([row[0] for row in results])
        
        # Generate response using Ollama
        try:
            response = requests.post(
                f"{OLLAMA_URL}/api/generate",
                json={
                    "model": OLLAMA_MODEL,
                    "prompt": (
                        f"You are a helpful NRI Banking Assistant. Be concise and friendly.\n"
                        f"Context from documents:\n{context}\n\n"
                        f"User: {query}\n"
                        f"Assistant:"
                    ),
                    "stream": False,
                },
                timeout=600,
            )
            response.raise_for_status()

            # Try to parse several possible Ollama response shapes to be resilient
            try:
                resp_json = response.json()
            except ValueError:
                # Not JSON - fall back to raw text
                resp_text = response.text or ""
                answer = resp_text.strip()
            else:
                logger.info(f"Ollama response JSON: {resp_json}")
                # Common keys: 'response', 'text', 'choices' (with content/text/message), 'result'
                answer = ""
                if isinstance(resp_json, dict):
                    if "response" in resp_json and resp_json.get("response"):
                        answer = resp_json.get("response")
                    elif "text" in resp_json and resp_json.get("text"):
                        answer = resp_json.get("text")
                    elif "result" in resp_json and resp_json.get("result"):
                        answer = resp_json.get("result")
                    elif "choices" in resp_json and isinstance(resp_json.get("choices"), list) and len(resp_json.get("choices")) > 0:
                        first = resp_json.get("choices")[0]
                        if isinstance(first, dict):
                            # try several possible fields
                            answer = (
                                first.get("content")
                                or first.get("text")
                                or first.get("message")
                                or first.get("response")
                                or ""
                            )
                        else:
                            answer = str(first)
                    else:
                        # As a last resort, stringify the whole response
                        answer = str(resp_json)
                else:
                    # Non-dict JSON (list/string)
                    answer = str(resp_json)

                if isinstance(answer, (list, dict)):
                    answer = str(answer)

                answer = (answer or "").strip()
                # If answer is empty after parsing, include the raw body for debugging
                if not answer:
                    answer = response.text or ""
        except Exception as e:
            answer = f"Error generating response: {str(e)}"
        
        # Save to chat history
        chat_entry = ChatHistory(
            user_query=query,
            answer=answer
        )
        db.add(chat_entry)
        db.commit()
        
        return {"answer": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()